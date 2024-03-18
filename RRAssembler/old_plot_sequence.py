"""
Make a drawing of the layout of a mapped contig. Run after the break_contigs_into_svs script
"""
import matplotlib.pyplot as plt
import pandas as pd
import sys
import numpy as np
from scipy import stats
import seaborn as sns
from collections import defaultdict
from matplotlib.pyplot import cm
import glob
import os
from collections import deque
from subprocess import call

import break_contigs_into_svs

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE


def get_underline(u, sim, contig):

    for s, e in sim:
        for i in range(s, e):
            if u[i] == "_":
                u[i] = "="  # Double underline
            elif u[i] == " ":
                u[i] = "_"
            elif u[i] == "=":
                u[i] = "B"
    return u


def new_line(current, current_under, blocks, under, keys, current_key, info_line):

    current += "\n"
    current_under += "\n"

    blocks.append(current)
    under.append(current_under)

    keys.append(current_key)
    keys.append("grey")

    blocks.append(info_line + "\n")
    under.append(" " * len(info_line) + "\n")

    current = ""  # Start a new current line, keep the same current key
    current_under = ""
    info_line = ""

    return current, current_under, blocks, under, keys, current_key, info_line


def contig_to_segments(contig, df):
    """
    Split the contig into a run of segments which can be stitched together. Different segments have different styles
    in the document. Buggy but gets the job done!
    :param contig:
    :param df:
    :return: list of segments with a code for the different styles
    """

    annotations = np.array(["u"] * len(contig))

    # Flesh out the aligned segments
    for s, e in zip(df.Q_start, df.Q_end):
        annotations[int(s): int(e)] = "a"

    # Add homology
    for i in range(len(df.Q_start) - 1):
        end1 = int(df.Q_end.iloc[i])
        start2 = int(df.Q_start.iloc[i+1])
        if start2 < end1:
            annotations[start2: end1] = "h"

    # Get the breakpoints in terms of the start positions on the contig
    breakpoints = zip(df.Q_start, zip(df.T_name, df.T_start, df.T_end))
    bp = set([i[0] for i in breakpoints])
    bp_ends = set(list(df.Q_end))

    #
    # Underline annotations
    underlines = [" "] * len(contig)
    THRESH = 0.05
    simularityA_se = [(int(i), int(j)) for i, j, P in zip(df.A_start, df.A_end, df.Pairwise_aln_prob)[:-1] if float(P) < THRESH]  # Last entry is empty!
    simularityB_se = [(int(i), int(j)) for i, j, P in zip(df.B_start, df.B_end, df.Pairwise_aln_prob)[:-1] if float(P) < THRESH]

    underlines = get_underline(underlines, simularityA_se, contig)
    underlines = get_underline(underlines, simularityB_se, contig)


    E_vals = iter(list(df.E))
    strands = iter(list(df.strand))
    pvals = iter([float(k) if float(k) < THRESH else 1 for k in list(df.Pairwise_aln_prob)])

    # Split list into contiguous blocks
    blocks = []
    keys = []
    current = ""
    current_key = annotations[0]

    under = []  # Same as blocks but shows ehere to draw underlines
    current_under = ""
    letters = 0
    position = 0

    pos_info = deque([])  # Contains the strand E value etc information, pop left items and add to info line
    # enables wrapping of strings over many lines
    info_line = ""

    for index, (k, base, undies) in enumerate(zip(annotations, contig, underlines)):

        # Add one letter at a time to both rows
        if letters == 89:
            current, current_under, blocks, under, keys, current_key, info_line = new_line(
                current, current_under, blocks, under, keys, current_key, info_line
            )

            letters = 0

        #
        # Deal with the info line
        if len(pos_info) == 0:
            info_line += " "
        else:
            item = pos_info.popleft()
            info_line += item

        if index in bp_ends:  # Add a | at the end of the alignment
            current += "|"

            current_under += " "
            # info_line += " "
            pos_info.append(" ")  # Dont add a space in the middle of some information
            letters += 1

        if index in bp:
            if len(current) > 0 and current[-1] != "|":
                current += "|"
                current_under += " "
                letters += 1

            # Add some information to the next line
            to_add = [i for i in breakpoints if i[0] == position][0]

            E = '{:0.0e}'.format(float(next(E_vals)))
            strand = next(strands)

            probability = float(next(pvals))
            if probability < THRESH:
                pval = '{:0.0e}'.format(probability)
                chrom = ">{}:{}-{} {} E={} p={} ".format(to_add[1][0], int(to_add[1][1]), int(to_add[1][2]), strand, E, pval)
            else:
                chrom = ">{}:{}-{} {} E={} ".format(to_add[1][0], int(to_add[1][1]), int(to_add[1][2]), strand, E)

            # Add to the deque
            for item in chrom:
                pos_info.append(item)

        # Check again if the line has reached the end
        if letters == 89:
            current, current_under, blocks, under, keys, current_key, info_line = new_line(
                current, current_under, blocks, under, keys, current_key, info_line
            )

            letters = 0

        # Add bases to the current block or reset the block
        if k == current_key:
            current += base
            current_under += undies
        else:
            blocks.append(current)
            under.append(current_under)

            keys.append(current_key)
            current = base
            current_under = undies

            current_key = k

        letters += 1
        position += 1


    while len(pos_info) > 0:
        current += pos_info.popleft()

    blocks.append(current + "\n")
    under.append(current_under)
    keys.append(current_key)

    # Add the last info line if it exists
    blocks.append(info_line)
    under.append(" "*len(info_line))
    keys.append("grey")

    return blocks, keys, under


def plot_alignments(df, contig, name):

    if not os.path.exists("annotated_contigs/{}".format(name.split("_")[0])):
        os.mkdir("annotated_contigs/{}".format(name.split("_")[0]))


    blocks, keys, under = contig_to_segments(contig, df)

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Courier'
    font.size = Pt(8)


    p = document.add_paragraph()
    run = p.add_run(name)
    run.font.size = Pt(10)
    run.font.name = "sans"
    run.font.bold = True

    # document.add_heading(name[1], level=1)
    p = document.add_paragraph()
    p.style = document.styles["Normal"]

    unmapped = (65, 105, 225)
    mapped = (0, 0, 0)
    homology = (220, 20, 60)
    grey = (150, 150, 150)

    for b, k, underline in zip(blocks, keys, under):
        for char, un in zip(b, underline):  # This is a bit of kluge to fix making the | characters all black instead of red and blue

            run = p.add_run(char)

            if char == "|":
                run.font.color.rgb = RGBColor(*mapped)

            elif k == "h":
                run.font.color.rgb = RGBColor(*homology)

            elif k == "u":
                run.font.color.rgb = RGBColor(*unmapped)

            elif k == "grey":
                run.font.color.rgb = RGBColor(*grey)

            if un == "_":
                run.font.underline = WD_UNDERLINE.SINGLE
            elif un == "=":
                run.font.underline = WD_UNDERLINE.DOUBLE
            elif un == "B":
                run.font.underline = WD_UNDERLINE.DOUBLE
                run.font.bold = True
        # else:
        #     run.font.color.rgb = RGBColor(*mapped)
            # run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            # run.font.underline = WD_UNDERLINE.WAVY
            # run.font.outline = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True

    document.add_page_break()
    document.save("annotated_contigs/{}/contig.{}.{}.docx".format(name.split("_")[0], len(df), name))


def process_alignments():

    samps = pd.read_csv("out_data/alignments.all.updated_psl.csv").groupby(["contig_names"])

    #call("tantan all_contigs.all.fa > out_data/all_contigs.all.tantan.fa", shell=True)

    conts = break_contigs_into_svs.load_contigs("all_contigs.all.fa")

    count = 0
    alns = 0
    for k, df in samps:
        # if "NODE_36_length_263_cov_0.578947" not in k:
        #     continue
        print(k, len(df))
        plot_alignments(df, conts[k], k)#"{}_{}".format(*k)].upper(), k)  # Ignore lowercase letters which are simple repeats
        # if count > 10:
        #     break
        count += 1
        alns += len(df)


    print(count, "contigs processed, ", alns, "alignments processed.")



if __name__ == "__main__":
    process_alignments()
